"""Rendering: profile dict -> .docx / .html / .pdf.

Clean modern minimalist: single column, Calibri body, bold ALL-CAPS section
headers with a bottom border rule, generous whitespace, no color, no tables for
layout. ATS-safe.

All three formats are derived from the same profile dict so they stay visually
consistent. PDF is produced by Playwright printing the HTML render.
"""

from __future__ import annotations

import html as _html
import re
from pathlib import Path
from typing import Any


def _collapse(text: str) -> str:
    """Collapse whitespace runs (including soft line breaks from YAML's `|`
    literal blocks) to single spaces. Use for single-paragraph inline fields:
    bullet text, blurbs, headings, etc."""
    return re.sub(r"\s+", " ", text or "").strip()


def _paragraphs(text: str) -> list[str]:
    """Split a multi-paragraph text on blank lines, collapsing soft line
    breaks within each paragraph. Returns a list of cleaned paragraphs."""
    text = (text or "").strip()
    if not text:
        return []
    parts = re.split(r"\n\s*\n", text)
    return [_collapse(p) for p in parts if p.strip()]

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor


BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(22)
HEADLINE_SIZE = Pt(11)
CONTACT_SIZE = Pt(9.5)
SECTION_SIZE = Pt(11.5)
ROLE_SIZE = Pt(11)


# ----- low-level helpers --------------------------------------------------

def _set_run(run, *, size=None, bold=False, italic=False, color=None,
             all_caps=False, font=BODY_FONT):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if size is not None:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if all_caps:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "true")
        rpr.append(caps)


def _set_paragraph_spacing(p, *, before=0, after=0, line=None, line_rule="auto"):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _add_bottom_border(paragraph, *, size_eighths=8, color="000000"):
    """Add a bottom rule under a paragraph (used for section headers)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))   # 8ths of a point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    # replace any existing bottom border
    existing = pbdr.find(qn("w:bottom"))
    if existing is not None:
        pbdr.remove(existing)
    pbdr.append(bottom)


def _add_tab_stops(paragraph, right_tab_inches):
    """Add a right-aligned tab stop, used for right-justified dates."""
    pf = paragraph.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(right_tab_inches),
                              alignment=WD_ALIGN_PARAGRAPH.RIGHT)


# ----- structural blocks --------------------------------------------------

def _add_name(doc: Document, name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run(name.upper())
    _set_run(run, size=NAME_SIZE, bold=True)
    rpr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "40")  # letter-spacing 40 twips
    rpr.append(spacing)


def _add_headline(doc: Document, headline: str):
    if not headline:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run(headline)
    _set_run(run, size=HEADLINE_SIZE, italic=True, color=RGBColor(0x33, 0x33, 0x33))


def _add_contact(doc: Document, identity: dict):
    parts = []
    for key in ("location", "email", "phone", "website", "linkedin", "github"):
        val = (identity.get(key) or "").strip()
        if not val:
            continue
        # Strip protocol for display
        display = val.replace("https://", "").replace("http://", "")
        parts.append(display)
    if not parts:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=10)
    run = p.add_run("  ·  ".join(parts))
    _set_run(run, size=CONTACT_SIZE, color=RGBColor(0x33, 0x33, 0x33))


def _add_section_header(doc: Document, label: str):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=10, after=4)
    run = p.add_run(label.upper())
    _set_run(run, size=SECTION_SIZE, bold=True, all_caps=True)
    rpr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")
    rpr.append(spacing)
    _add_bottom_border(p, size_eighths=6)


def _add_summary(doc: Document, summary: str):
    if not summary or summary.strip().startswith("TODO"):
        return
    paragraphs = _paragraphs(summary)
    for i, ptext in enumerate(paragraphs):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0 if i == 0 else 4, after=4, line=1.15)
        run = p.add_run(ptext)
        _set_run(run, size=BODY_SIZE)


def _year_only(s: str) -> str:
    """Normalize dates to YYYY (or 'Present'). Accepts '', 'YYYY', 'YYYY-MM',
    'YYYY-MM-DD', 'Present', 'present'."""
    s = (s or "").strip()
    if not s:
        return ""
    if s.lower() == "present":
        return "Present"
    # Take the leading 4-digit year if present
    import re
    m = re.match(r"\d{4}", s)
    return m.group(0) if m else s


def _format_date_range(start: str, end: str) -> str:
    start_y = _year_only(start)
    end_y = _year_only(end) or "Present"
    if not start_y and end_y == "Present":
        return ""
    if not start_y:
        return end_y
    if start_y == end_y:
        return start_y
    return f"{start_y} – {end_y}"


def _add_experience_entry(doc: Document, role: dict, right_tab: float):
    company = (role.get("company") or "").strip()
    location = (role.get("location") or "").strip()
    title = (role.get("title") or "").strip()
    date_range = _format_date_range(role.get("start"), role.get("end"))

    # Line 1: Company · Location   <tab>   Dates
    line1 = doc.add_paragraph()
    _set_paragraph_spacing(line1, before=4, after=0)
    _add_tab_stops(line1, right_tab)
    left = company
    if location:
        left = f"{company}  ·  {location}" if company else location
    r = line1.add_run(left)
    _set_run(r, size=ROLE_SIZE, bold=True)
    if date_range:
        r2 = line1.add_run(f"\t{date_range}")
        _set_run(r2, size=BODY_SIZE, color=RGBColor(0x33, 0x33, 0x33))

    # Line 2: Title (italic)
    if title:
        line2 = doc.add_paragraph()
        _set_paragraph_spacing(line2, before=0, after=2)
        r = line2.add_run(title)
        _set_run(r, size=BODY_SIZE, italic=True, color=RGBColor(0x33, 0x33, 0x33))

    # Blurb
    blurb = _collapse(role.get("blurb") or "")
    if blurb:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=2, line=1.15)
        r = p.add_run(blurb)
        _set_run(r, size=BODY_SIZE)

    # Bullets
    for bullet in role.get("bullets", []) or []:
        if isinstance(bullet, str):
            text = _collapse(bullet)
        else:
            text = _collapse(bullet.get("text") or "")
            if not text:
                continue
            if not bullet.get("general", True):
                continue
        if not text:
            continue
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_spacing(p, before=0, after=2, line=1.15)
        for run in p.runs:
            _set_run(run, size=BODY_SIZE)
        # The List Bullet style adds the bullet itself; add the text:
        r = p.add_run(text)
        _set_run(r, size=BODY_SIZE)


def _add_skills(doc: Document, skill_groups: list):
    any_rendered = False
    for group in skill_groups or []:
        items = [i for i in (group.get("items") or []) if str(i).strip()]
        if not items:
            continue
        any_rendered = True
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=2, line=1.15)
        label = (group.get("group") or "").strip()
        if label:
            r = p.add_run(f"{label}:  ")
            _set_run(r, size=BODY_SIZE, bold=True)
        r = p.add_run(", ".join(str(i) for i in items))
        _set_run(r, size=BODY_SIZE)
    return any_rendered


def _add_education_entry(doc: Document, edu: dict, right_tab: float):
    institution = (edu.get("institution") or "").strip()
    location = (edu.get("location") or "").strip()
    date_range = _format_date_range(edu.get("start"), edu.get("end"))
    degree = (edu.get("degree") or "").strip()
    field = (edu.get("field") or "").strip()

    line1 = doc.add_paragraph()
    _set_paragraph_spacing(line1, before=4, after=0)
    _add_tab_stops(line1, right_tab)
    left = institution
    if location:
        left = f"{institution}  ·  {location}" if institution else location
    r = line1.add_run(left)
    _set_run(r, size=ROLE_SIZE, bold=True)
    if date_range:
        r2 = line1.add_run(f"\t{date_range}")
        _set_run(r2, size=BODY_SIZE, color=RGBColor(0x33, 0x33, 0x33))

    degree_line = ", ".join(x for x in [degree, field] if x)
    if degree_line:
        line2 = doc.add_paragraph()
        _set_paragraph_spacing(line2, before=0, after=2)
        r = line2.add_run(degree_line)
        _set_run(r, size=BODY_SIZE, italic=True, color=RGBColor(0x33, 0x33, 0x33))

    extras = []
    if (gpa := (edu.get("gpa") or "").strip()):
        extras.append(f"GPA {gpa}")
    if honors := (edu.get("honors") or []):
        extras.append(", ".join(honors))
    if extras:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=2)
        r = p.add_run("  ·  ".join(extras))
        _set_run(r, size=BODY_SIZE)


def _add_cert_entry(doc: Document, cert: dict, right_tab: float):
    name = (cert.get("name") or "").strip()
    issuer = (cert.get("issuer") or "").strip()
    date = (cert.get("date") or "").strip()
    if not name:
        return
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=2, after=2)
    _add_tab_stops(p, right_tab)
    text = name + (f"  ·  {issuer}" if issuer else "")
    r = p.add_run(text)
    _set_run(r, size=BODY_SIZE)
    if date:
        r2 = p.add_run(f"\t{date}")
        _set_run(r2, size=BODY_SIZE, color=RGBColor(0x33, 0x33, 0x33))


# ----- main entry --------------------------------------------------------

def _set_page(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    # Default Normal style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE


def render(profile: dict, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    doc = Document()
    _set_page(doc)

    identity = profile.get("identity") or {}
    _add_name(doc, identity.get("name") or "")
    _add_headline(doc, identity.get("headline") or "")
    _add_contact(doc, identity)

    # Right tab position (approximate: page width 8.5 - 0.7*2 = 7.1)
    right_tab = 7.1

    summary = (profile.get("summary") or "").strip()
    if summary and not summary.startswith("TODO"):
        _add_section_header(doc, "Summary")
        _add_summary(doc, summary)

    experience = profile.get("experience") or []
    if any((e.get("company") or e.get("title")) for e in experience):
        _add_section_header(doc, "Experience")
        for role in experience:
            if not (role.get("company") or role.get("title")):
                continue
            _add_experience_entry(doc, role, right_tab)

    skill_groups = profile.get("skills") or []
    if any((g.get("items") or []) for g in skill_groups):
        _add_section_header(doc, "Skills")
        _add_skills(doc, skill_groups)

    education = profile.get("education") or []
    if any((e.get("institution") or e.get("degree")) for e in education):
        _add_section_header(doc, "Education")
        for edu in education:
            if not (edu.get("institution") or edu.get("degree")):
                continue
            _add_education_entry(doc, edu, right_tab)

    certs = profile.get("certifications") or []
    if any(c.get("name") for c in certs):
        _add_section_header(doc, "Certifications")
        for cert in certs:
            _add_cert_entry(doc, cert, right_tab)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ============================================================================
# HTML renderer — visually mirrors the docx; doubles as the PDF source.
# ============================================================================

_HTML_STYLE = """
:root { color-scheme: light; }
html, body { margin: 0; padding: 0; background: #f3f3f3; }
@page { size: Letter; margin: 0.5in 0.7in; }
body {
  font-family: Calibri, "Helvetica Neue", Helvetica, Arial, sans-serif;
  font-size: 10.5pt;
  color: #1a1a1a;
  line-height: 1.25;
}
.page {
  background: #fff;
  width: 7.1in;
  margin: 16px auto;
  padding: 0.5in 0.7in;
  box-sizing: border-box;
  box-shadow: 0 1px 4px rgba(0,0,0,0.08);
}
@media print {
  body { background: #fff; }
  .page { margin: 0; padding: 0; width: auto; box-shadow: none; }
}
.name {
  text-align: center;
  font-size: 22pt;
  font-weight: 700;
  letter-spacing: 0.08em;
  margin: 0;
}
.headline {
  text-align: center;
  font-size: 11pt;
  font-style: italic;
  color: #333;
  margin: 2pt 0 2pt;
}
.contact {
  text-align: center;
  font-size: 9.5pt;
  color: #333;
  margin: 0 0 10pt;
}
.section-h {
  font-size: 11.5pt;
  font-weight: 700;
  text-transform: uppercase;
  letter-spacing: 0.05em;
  border-bottom: 0.75pt solid #000;
  padding-bottom: 1pt;
  margin: 10pt 0 4pt;
}
.summary { margin: 0 0 4pt; line-height: 1.3; }
.role { margin-top: 4pt; }
.role-line1 {
  display: flex;
  justify-content: space-between;
  font-size: 11pt;
  font-weight: 700;
  margin: 0;
}
.role-line1 .dates { font-weight: 400; color: #333; font-size: 10.5pt; }
.role-title {
  font-style: italic;
  color: #333;
  margin: 0 0 2pt;
}
.role-blurb { margin: 0 0 2pt; }
ul.bullets { margin: 2pt 0 4pt 18pt; padding: 0; }
ul.bullets li { margin: 0 0 2pt; line-height: 1.3; }
.skill-row { margin: 2pt 0; line-height: 1.3; }
.skill-group { font-weight: 700; }
.edu-line1 {
  display: flex;
  justify-content: space-between;
  font-size: 11pt;
  font-weight: 700;
  margin: 4pt 0 0;
}
.edu-line1 .dates { font-weight: 400; color: #333; font-size: 10.5pt; }
.edu-degree { font-style: italic; color: #333; margin: 0 0 2pt; }
.edu-extras { margin: 0 0 2pt; }
.cert {
  display: flex;
  justify-content: space-between;
  margin: 2pt 0;
}
.cert .dates { color: #333; }
"""


def _esc(s: Any) -> str:
    return _html.escape(str(s or ""))


def _contact_html(identity: dict) -> str:
    parts = []
    for key in ("location", "email", "phone", "website", "linkedin", "github"):
        val = (identity.get(key) or "").strip()
        if val:
            display = val.replace("https://", "").replace("http://", "")
            parts.append(_esc(display))
    return "  ·  ".join(parts)


def _ed_attrs(path: str | None) -> str:
    """contenteditable+data-path attributes, only when path is provided (edit mode)."""
    if not path:
        return ""
    return f' contenteditable="true" data-path="{path}" spellcheck="true"'


def _del_btn(path: str | None, label: str = "✕", title: str = "Delete") -> str:
    """Small inline delete button. Renders only in editable mode."""
    if not path:
        return ""
    return (
        f'<button type="button" class="del-btn" data-del-path="{path}" '
        f'title="{title}" tabindex="-1">{label}</button>'
    )


def _role_html_paths(role: dict, idx: int, editable: bool) -> str:
    """role rendering with optional edit attributes."""
    company = (role.get("company") or "").strip()
    location = (role.get("location") or "").strip()
    title = (role.get("title") or "").strip()
    date_range = _format_date_range(role.get("start"), role.get("end"))
    left = company
    if location:
        left = f"{company}  ·  {location}" if company else location

    pre = f"experience.{idx}" if editable else None
    role_del = _del_btn(pre, label="✕ role", title="Remove this entire role")
    out = [f'<div class="role">{role_del}']
    out.append(f'<p class="role-line1"><span{_ed_attrs(pre + ".company") if pre else ""}>{_esc(left)}</span>')
    if date_range:
        out.append(f'<span class="dates">{_esc(date_range)}</span>')
    out.append('</p>')
    if title:
        out.append(f'<p class="role-title"{_ed_attrs(pre + ".title") if pre else ""}>{_esc(title)}</p>')
    blurb = _collapse(role.get("blurb") or "")
    if blurb:
        out.append(f'<p class="role-blurb"{_ed_attrs(pre + ".blurb") if pre else ""}>{_esc(blurb)}</p>')

    bullets_html = []
    # Iterate full list with original indices so post-skip data-paths stay stable.
    for j, b in enumerate(role.get("bullets", []) or []):
        if isinstance(b, str):
            text = _collapse(b)
            bpath = f"{pre}.bullets.{j}" if pre else None
        else:
            if not b.get("general", True):
                continue
            text = _collapse(b.get("text") or "")
            bpath = f"{pre}.bullets.{j}.text" if pre else None
        if text:
            # × deletes the WHOLE bullet entry, so address the parent index.
            del_path = f"{pre}.bullets.{j}" if pre else None
            bullets_html.append(
                f'<li{_ed_attrs(bpath)}>{_del_btn(del_path)}{_esc(text)}</li>'
            )
    if bullets_html:
        out.append('<ul class="bullets">' + "".join(bullets_html) + '</ul>')
    out.append('</div>')
    return "".join(out)


def _section_h(label: str, del_path: str | None) -> str:
    return (
        f'<h2 class="section-h">{_esc(label)}'
        f'{_del_btn(del_path, label="✕ remove section", title=f"Remove the {label} section")}'
        f'</h2>'
    )


def render_html(profile: dict, *, editable: bool = False,
                save_url: str | None = None,
                delete_url: str | None = None) -> str:
    """Render the resume as a standalone HTML document.

    When editable=True, every text field carries contenteditable + a data-path
    attribute, and a small inline script POSTs blur events to `save_url` as
    JSON `{path, text}`. The visual layout is identical to the read-only view.
    """
    identity = profile.get("identity") or {}
    name = identity.get("name") or ""
    headline = (identity.get("headline") or "").strip()
    contact = _contact_html(identity)

    body_parts = [
        f'<h1 class="name"{_ed_attrs("identity.name") if editable else ""}>{_esc(name.upper())}</h1>'
    ]
    if headline:
        body_parts.append(
            f'<p class="headline"{_ed_attrs("identity.headline") if editable else ""}>{_esc(headline)}</p>'
        )
    if contact:
        # contact is composed of multiple sub-fields; not inline-editable here
        # (use the metadata editor on the job page). Keep displayed as one line.
        body_parts.append(f'<p class="contact">{contact}</p>')

    summary_raw = (profile.get("summary") or "").strip()
    if summary_raw and not summary_raw.startswith("TODO"):
        body_parts.append(_section_h("Summary", "summary" if editable else None))
        paragraphs = _paragraphs(summary_raw)
        # In edit mode keep summary as one editable paragraph so the user can
        # type freely; on save we collapse to whatever YAML form survives a
        # round-trip through render.
        if editable:
            joined = "\n\n".join(paragraphs)
            body_parts.append(
                f'<p class="summary"{_ed_attrs("summary")}>{_esc(joined)}</p>'
            )
        else:
            for ptext in paragraphs:
                body_parts.append(f'<p class="summary">{_esc(ptext)}</p>')

    experience = profile.get("experience") or []
    visible_indices = [
        i for i, r in enumerate(experience)
        if (r.get("company") or r.get("title")) and r.get("general", True) is not False
    ]
    if visible_indices:
        body_parts.append(_section_h("Experience", "experience" if editable else None))
        for i in visible_indices:
            body_parts.append(_role_html_paths(experience[i], i, editable))

    skill_groups = profile.get("skills") or []
    if any((g.get("items") or []) for g in skill_groups):
        body_parts.append(_section_h("Skills", "skills" if editable else None))
        for gi, g in enumerate(skill_groups):
            items = [str(i) for i in (g.get("items") or []) if str(i).strip()]
            if not items:
                continue
            label = (g.get("group") or "").strip()
            lbl_html = (
                f'<span class="skill-group"{_ed_attrs(f"skills.{gi}.group") if editable else ""}>{_esc(label)}:</span> '
                if label else ""
            )
            items_html = (
                f'<span{_ed_attrs(f"skills.{gi}.items") if editable else ""}>{_esc(", ".join(items))}</span>'
            )
            body_parts.append(f'<p class="skill-row">{lbl_html}{items_html}</p>')

    education = profile.get("education") or []
    visible_edu = [(i, e) for i, e in enumerate(education) if (e.get("institution") or e.get("degree"))]
    if visible_edu:
        body_parts.append(_section_h("Education", "education" if editable else None))
        for i, e in visible_edu:
            body_parts.append(_edu_html_ed(e, i, editable))

    certs = profile.get("certifications") or []
    visible_certs = [(i, c) for i, c in enumerate(certs) if c.get("name")]
    if visible_certs:
        body_parts.append(_section_h("Certifications", "certifications" if editable else None))
        for i, c in visible_certs:
            body_parts.append(_cert_html_ed(c, i, editable))

    editor_script = ""
    if editable and save_url:
        editor_script = (_EDIT_SCRIPT
                         .replace("__SAVE_URL__", save_url)
                         .replace("__DELETE_URL__", delete_url or ""))

    return (
        f'<!doctype html><html><head><meta charset="utf-8">'
        f'<title>{_esc(name)} — Resume{" (editing)" if editable else ""}</title>'
        f'<style>{_HTML_STYLE}{_EDIT_STYLE if editable else ""}</style></head>'
        f'<body><div class="page">{"".join(body_parts)}</div>'
        f'{editor_script}'
        f'</body></html>'
    )


def _edu_html_ed(edu: dict, idx: int, editable: bool) -> str:
    institution = (edu.get("institution") or "").strip()
    location = (edu.get("location") or "").strip()
    date_range = _format_date_range(edu.get("start"), edu.get("end"))
    degree = (edu.get("degree") or "").strip()
    field = (edu.get("field") or "").strip()
    left = institution
    if location:
        left = f"{institution}  ·  {location}" if institution else location
    pre = f"education.{idx}" if editable else None

    out = []
    out.append(
        f'<p class="edu-line1"><span{_ed_attrs(pre + ".institution") if pre else ""}>{_esc(left)}</span>'
    )
    if date_range:
        out.append(f'<span class="dates">{_esc(date_range)}</span>')
    out.append('</p>')
    degree_line = ", ".join(x for x in [degree, field] if x)
    if degree_line:
        out.append(
            f'<p class="edu-degree"{_ed_attrs(pre + ".field") if pre else ""}>{_esc(degree_line)}</p>'
        )

    extras = []
    if (gpa := (edu.get("gpa") or "").strip()):
        extras.append(f"GPA {gpa}")
    if honors := (edu.get("honors") or []):
        extras.append(", ".join(honors))
    if extras:
        out.append(f'<p class="edu-extras">{_esc("  ·  ".join(extras))}</p>')
    return "".join(out)


def _cert_html_ed(cert: dict, idx: int, editable: bool) -> str:
    name = (cert.get("name") or "").strip()
    issuer = (cert.get("issuer") or "").strip()
    date = (cert.get("date") or "").strip()
    if not name:
        return ""
    pre = f"certifications.{idx}" if editable else None
    label = name + (f"  ·  {issuer}" if issuer else "")
    out = [
        f'<p class="cert"><span{_ed_attrs(pre + ".name") if pre else ""}>{_esc(label)}</span>'
    ]
    if date:
        out.append(f'<span class="dates">{_esc(date)}</span>')
    out.append('</p>')
    return "".join(out)


_EDIT_STYLE = """
[contenteditable] {
  outline: none;
  border-radius: 3px;
  transition: background 0.1s;
}
[contenteditable]:hover {
  background: #fffbe6;
  box-shadow: 0 0 0 1px #f0d878 inset;
}
[contenteditable]:focus {
  background: #fff8d4;
  box-shadow: 0 0 0 2px #d4a017 inset;
}
[contenteditable].saved {
  background: #e6f7ec !important;
  transition: background 1.5s;
}
[contenteditable].save-error {
  background: #ffebee !important;
  box-shadow: 0 0 0 2px #b00020 inset !important;
}
#__edit-status {
  position: fixed; bottom: 16px; right: 16px;
  background: #1a1a1a; color: #fff; padding: 6px 12px;
  border-radius: 6px; font-family: system-ui, sans-serif; font-size: 12px;
  opacity: 0; transition: opacity 0.2s;
  z-index: 9999;
}
#__edit-status.show { opacity: 0.92; }
.edit-banner {
  background: #fff5e6; border: 1px solid #f0c878; padding: 8px 12px;
  border-radius: 6px; font-family: system-ui, sans-serif; font-size: 13px;
  margin: 0 0 12px;
}
.del-btn {
  display: inline-block;
  margin-left: 6px;
  padding: 0 6px;
  font-family: system-ui, sans-serif; font-size: 11px;
  background: #fff; color: #b00020;
  border: 1px solid #e0c0c0; border-radius: 4px;
  cursor: pointer; opacity: 0.4;
  transition: opacity 0.1s;
  vertical-align: middle;
}
.del-btn:hover { opacity: 1; background: #ffebee; }
.section-h .del-btn { font-size: 10pt; font-weight: 400; text-transform: none; }
li .del-btn { margin-right: 4px; margin-left: -22px; }
.role { position: relative; }
.role > .del-btn { position: absolute; top: 4px; right: 0; }
"""


_EDIT_SCRIPT = """
<div id="__edit-status">saved</div>
<script>
(function() {
  const SAVE_URL = "__SAVE_URL__";
  const DELETE_URL = "__DELETE_URL__";
  const statusEl = document.getElementById("__edit-status");
  function flash(msg) {
    statusEl.textContent = msg;
    statusEl.classList.add("show");
    setTimeout(() => statusEl.classList.remove("show"), 1200);
  }
  async function save(el) {
    const path = el.dataset.path;
    const text = el.innerText.replace(/\\u00A0/g, " ").trim();
    el.classList.remove("saved", "save-error");
    try {
      const r = await fetch(SAVE_URL, {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({path, text})
      });
      if (!r.ok) {
        const err = await r.json().catch(() => ({error: r.statusText}));
        el.classList.add("save-error");
        flash("error: " + (err.error || r.statusText));
        return;
      }
      el.classList.add("saved");
      flash("saved");
    } catch (e) {
      el.classList.add("save-error");
      flash("network error");
    }
  }
  async function del(path, confirmMsg) {
    if (!DELETE_URL) return;
    if (confirmMsg && !confirm(confirmMsg)) return;
    flash("deleting…");
    try {
      const r = await fetch(DELETE_URL, {
        method: "POST",
        headers: {"Content-Type": "application/json"},
        body: JSON.stringify({path})
      });
      if (!r.ok) {
        const err = await r.json().catch(() => ({error: r.statusText}));
        flash("error: " + (err.error || r.statusText));
        return;
      }
      // Reload so indexes/data-paths re-derive from the updated YAML.
      location.reload();
    } catch (e) {
      flash("network error");
    }
  }
  document.querySelectorAll("[contenteditable][data-path]").forEach(el => {
    let originalText = el.innerText;
    el.addEventListener("focus", () => { originalText = el.innerText; });
    el.addEventListener("blur", () => {
      const newText = el.innerText.replace(/\\u00A0/g, " ").trim();
      if (newText === originalText.trim()) return;
      save(el);
    });
    // Plain-text paste (strip formatting from clipboard)
    el.addEventListener("paste", e => {
      e.preventDefault();
      const t = (e.clipboardData || window.clipboardData).getData("text/plain");
      document.execCommand("insertText", false, t);
    });
  });
  document.querySelectorAll(".del-btn[data-del-path]").forEach(btn => {
    btn.addEventListener("click", e => {
      e.preventDefault();
      e.stopPropagation();
      const path = btn.dataset.delPath;
      // Confirm only for big deletes (whole sections / whole roles)
      const big = !path.includes(".bullets.") || path.split(".").length <= 1;
      del(path, big ? ("Remove " + path + "?") : null);
    });
  });
  // Banner at top
  const banner = document.createElement("div");
  banner.className = "edit-banner";
  banner.innerHTML = "<strong>Edit mode.</strong> Click any text to edit it. <strong>Changes save automatically the moment you click anywhere else</strong> — there is no Save button (and you do not need one). Watch for the green flash + the &ldquo;saved&rdquo; indicator in the bottom-right corner. Click <span style=\\"color:#b00020\\">✕</span> to remove a bullet, role, or whole section. When you are done editing, go back to the job page and click <em>Build resume</em> to refresh the .docx / .pdf for download.";
  document.body.querySelector(".page").prepend(banner);
})();
</script>
"""


# ============================================================================
# PDF renderer — Playwright prints the HTML.
# ============================================================================

def render_pdf(html: str, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html, wait_until="domcontentloaded")
        page.pdf(
            path=str(output_path),
            format="Letter",
            margin={"top": "0.5in", "bottom": "0.5in",
                    "left": "0.7in", "right": "0.7in"},
            print_background=True,
            prefer_css_page_size=True,
        )
        browser.close()
    return output_path


def render_all(profile: dict, base_path: str | Path, *, strict: bool = True) -> dict:
    """Render docx, html, pdf side by side. Returns paths + lint findings.

    If `strict` (default), any lint ERROR raises LintError BEFORE writing files.
    Warnings (AI-tells) never block; they're returned for the UX to surface.
    """
    from lint_resume import lint, errors, warnings, LintError
    findings = lint(profile)
    if strict and (errs := errors(findings)):
        raise LintError(errs)

    base = Path(base_path)
    docx_path = base.with_suffix(".docx")
    html_path = base.with_suffix(".html")
    pdf_path = base.with_suffix(".pdf")

    render(profile, docx_path)
    html = render_html(profile)
    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(html, encoding="utf-8")
    render_pdf(html, pdf_path)

    return {
        "docx": docx_path, "html": html_path, "pdf": pdf_path,
        "warnings": warnings(findings),
        "errors": errors(findings),
    }
