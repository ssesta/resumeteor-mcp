"""Unit tests for scripts/lint_resume."""

import sys
from pathlib import Path

import pytest
import yaml

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "scripts"))


def _profile(**overrides) -> dict:
    base = {
        "identity": {"name": "X", "email": "x@y.z"},
        "summary": "Clean summary.",
        "experience": [{
            "company": "Acme",
            "title": "Engineer",
            "start": "2020",
            "end": "Present",
            "bullets": [{"text": "Did a thing.", "general": True}],
        }],
        "education": [{"institution": "U", "degree": "BS", "field": "CS"}],
        "skills": [{"group": "G", "items": ["x"]}],
    }
    base.update(overrides)
    return base


def test_clean_profile_has_no_findings():
    from lint_resume import lint
    assert lint(_profile()) == []


def test_html_entity_in_text_is_error():
    from lint_resume import lint, errors
    p = _profile()
    p["experience"][0]["bullets"][0]["text"] = "Worked with R&D &middot; great team."
    errs = errors(lint(p))
    assert any(e.code == "html_entity_leak" for e in errs)


def test_todo_marker_is_error():
    from lint_resume import lint, errors
    p = _profile()
    p["summary"] = "TODO fill this in"
    assert any(e.code == "todo_marker" for e in errors(lint(p)))


def test_template_syntax_is_error():
    from lint_resume import lint, errors
    p = _profile()
    p["experience"][0]["bullets"][0]["text"] = "Built {{ thing }} for clients."
    assert any(e.code == "raw_template_syntax" for e in errors(lint(p)))


def test_ai_tell_phrase_is_warning_not_error():
    from lint_resume import lint, errors, warnings
    p = _profile()
    p["experience"][0]["bullets"][0]["text"] = "Spearheaded a cross-functional initiative."
    findings = lint(p)
    assert errors(findings) == []
    warns = warnings(findings)
    assert any(w.code == "ai_tell" and "spearhead" in w.message.lower() for w in warns)


def test_ai_tell_in_short_field_is_ignored():
    """Don't flag a single keyword in something like institution name."""
    from lint_resume import lint, warnings
    p = _profile()
    p["education"][0]["institution"] = "Synergy College"
    assert warnings(lint(p)) == []


def test_build_with_entity_leak_raises(tmp_path):
    from build_resume import build_general
    import os
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    p["summary"] = "Used &middot; widely."
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from lint_resume import LintError
    with pytest.raises(LintError):
        build_general(root=tmp_path)


def test_dates_render_year_only_in_html(tmp_path):
    """End-to-end: profile with YYYY-MM start renders as YYYY in HTML."""
    import os
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    p["experience"][0]["start"] = "2020-03"
    p["experience"][0]["end"] = "2024-11"
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from build_resume import build_general
    out = build_general(root=tmp_path)
    html = out["html"].read_text()
    assert "2020 – 2024" in html
    assert "2020-03" not in html
    assert "2024-11" not in html


def test_yaml_pipe_block_summary_renders_as_single_docx_paragraph(tmp_path):
    """Soft line breaks from YAML's `|` literal-block syntax must collapse to
    spaces in the docx, not survive as in-paragraph newlines that wrap text
    at column width."""
    import os
    from docx import Document
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    # mimic the YAML `|` form — hard newlines mid-sentence
    p["summary"] = (
        "Senior post-sale technology leader with 25+ years scaling\n"
        "customer-success and technical-services organizations for\n"
        "enterprise and strategic accounts."
    )
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from build_resume import build_general
    out = build_general(root=tmp_path)
    doc = Document(str(out["docx"]))
    # Find the summary paragraph
    summary_paras = [pp for pp in doc.paragraphs
                     if pp.text.startswith("Senior post-sale")]
    assert summary_paras, "summary paragraph not found"
    # No newline characters should remain in any run of the summary
    for pp in summary_paras:
        for run in pp.runs:
            assert "\n" not in run.text, (
                f"summary still has in-paragraph newline: {run.text!r}")
    # Should appear as a continuous, single-spaced sentence
    assert "scaling customer-success" in summary_paras[0].text
    assert "  " not in summary_paras[0].text  # no double-space artifacts


def test_yaml_blank_lines_become_separate_summary_paragraphs(tmp_path):
    """A blank line in the source should produce a real paragraph break."""
    import os
    from docx import Document
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    p["summary"] = "First paragraph here.\n\nSecond paragraph here."
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from build_resume import build_general
    out = build_general(root=tmp_path)
    doc = Document(str(out["docx"]))
    texts = [pp.text for pp in doc.paragraphs]
    assert "First paragraph here." in texts
    assert "Second paragraph here." in texts


def test_bullet_text_collapses_soft_breaks(tmp_path):
    """Multi-line bullet text from YAML `|` should render as one continuous line."""
    import os
    from docx import Document
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    p["experience"][0]["bullets"][0]["text"] = (
        "Built and operated a thing across\nmultiple lines from the YAML source."
    )
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from build_resume import build_general
    out = build_general(root=tmp_path)
    doc = Document(str(out["docx"]))
    bullets = [pp.text for pp in doc.paragraphs if "thing" in pp.text]
    assert bullets, "bullet not found"
    assert "\n" not in bullets[0]
    assert "thing across multiple lines" in bullets[0]


def test_no_html_entities_in_rendered_html(tmp_path):
    """No &middot; should appear as literal text in the HTML output."""
    import os, re
    (tmp_path / "kitchen-sink").mkdir()
    (tmp_path / "general-resume").mkdir()
    p = _profile()
    p["experience"][0]["company"] = "Acme"
    p["experience"][0]["location"] = "NYC"
    (tmp_path / "kitchen-sink" / "profile.yaml").write_text(
        yaml.safe_dump(p, sort_keys=False)
    )
    os.environ["JOBSEARCH_ROOT"] = str(tmp_path)
    from build_resume import build_general
    out = build_general(root=tmp_path)
    html = out["html"].read_text()
    # Allow CSS entities (none expected) but no literal &middot; or &amp;middot;
    assert "&middot;" not in html
    assert "&amp;middot;" not in html
    # The real character should appear where company · location is rendered
    assert "Acme  ·  NYC" in html
